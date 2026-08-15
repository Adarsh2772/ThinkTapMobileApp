"""Generate ThinkTap Backend API Implementation Guide (Word)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    return h


def add_para(doc, text, bold=False, size=11, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for run in p.runs:
        set_run_font(run)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        set_run_font(run)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.15)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
    doc.add_paragraph()
    return table


def build():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Think Tap")
    set_run_font(r, size=28, bold=True, color=RGBColor(0x4B, 0x41, 0xE1))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Backend API Implementation Guide")
    set_run_font(r, size=20, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "Node.js + Express — Step-by-step guide for backend developers\n"
        "Version 1.0  |  Based on Think Tap (Expo) mobile codebase\n"
        "Audience: Backend engineers implementing cloud Auth, Ideas, Audio, and AI APIs"
    )
    set_run_font(r, size=11, color=RGBColor(0x6B, 0x72, 0x80))

    doc.add_paragraph()

    # 1. Overview
    add_heading(doc, "1. Product & Current State Overview", 1)
    add_para(
        doc,
        "Think Tap is a voice-first mobile idea-capture app (Expo / React Native) for Android and iOS. "
        "The user taps the mic (or says “Hey Think Tap”), records audio, and the app produces a titled, "
        "categorized, summarized idea with a full transcript and optional AI story rewrite.",
    )
    add_para(doc, "Core user loop", bold=True)
    add_numbered(doc, "Start recording (tap or wake word).")
    add_numbered(doc, "Stop recording (tap or voice “Stop”).")
    add_numbered(doc, "Processing screen: upload → transcribe → extract → summarize → save.")
    add_numbered(doc, "Idea detail: play audio, read summary & transcript, favorite.")
    add_numbered(doc, "Ideas list / Search: browse, filter by category, text search.")

    add_para(doc, "What exists today vs what you must build", bold=True)
    add_table(
        doc,
        ["Area", "Mobile app today", "Backend target"],
        [
            ["Auth", "Local guest session + AsyncStorage users (plaintext passwords)", "Real signup/login, hashed passwords, JWT"],
            ["Ideas", "Stored in AsyncStorage on device", "Postgres (or similar) per-user CRUD"],
            ["Audio", "Local file:// paths under documentDirectory", "Object storage + signed URLs"],
            ["AI / STT", "Device OS STT + local heuristics; optional STT proxy", "Server-side Whisper + LLM; keys never in app"],
            ["Settings", "Local AsyncStorage", "Optional /me/settings sync"],
            ["Wake word", "On-device only (no network)", "No backend APIs needed"],
            ["Existing server", "server/ Express proxy: /health, /api/transcribe, /api/enrich", "Extend into full API"],
        ],
    )
    add_para(
        doc,
        "Important: A thin Node proxy already lives in the repo at server/. Production keys must stay "
        "server-side only. Do not ship GROQ/OpenAI keys as EXPO_PUBLIC_* in production builds.",
    )

    # 2. Architecture
    add_heading(doc, "2. Recommended Architecture (Node.js)", 1)
    add_para(doc, "Recommended stack", bold=True)
    add_table(
        doc,
        ["Layer", "Recommendation", "Notes"],
        [
            ["Runtime", "Node.js 20+", "Match Expo project tooling"],
            ["Framework", "Express (extend existing server/) or NestJS", "Existing proxy is Express"],
            ["Database", "PostgreSQL (Supabase or self-hosted)", "README recommends Supabase for MVP"],
            ["ORM", "Prisma or Drizzle", "Typed models matching mobile Idea/User types"],
            ["Auth", "JWT access + refresh tokens (or Supabase Auth)", "Bearer Authorization header"],
            ["File storage", "S3 / Cloudflare R2 / Supabase Storage", "Private bucket; signed read URLs"],
            ["STT", "Groq whisper-large-v3 (cost) or OpenAI gpt-4o-transcribe", "Auto-detect language; omit language param"],
            ["LLM", "Groq llama-3.3-70b-versatile or OpenAI gpt-4o-mini", "JSON: title, category, summary, aiStory"],
            ["Validation", "Zod or Joi", "Validate bodies & multipart"],
            ["Security", "helmet, cors allowlist, rate-limit, TLS via reverse proxy", "Required for production"],
        ],
    )

    add_para(doc, "Suggested folder structure", bold=True)
    add_code(
        doc,
        "backend/\n"
        "├── src/\n"
        "│   ├── index.js                 # app bootstrap\n"
        "│   ├── config.js                # env, limits, models\n"
        "│   ├── middleware/\n"
        "│   │   ├── auth.js              # JWT verify\n"
        "│   │   ├── upload.js            # multer (25 MB)\n"
        "│   │   ├── errorHandler.js\n"
        "│   │   └── rateLimit.js\n"
        "│   ├── routes/\n"
        "│   │   ├── auth.routes.js\n"
        "│   │   ├── ideas.routes.js\n"
        "│   │   ├── audio.routes.js\n"
        "│   │   ├── ai.routes.js         # enrich / transcribe\n"
        "│   │   ├── settings.routes.js\n"
        "│   │   └── subscriptions.routes.js\n"
        "│   ├── services/\n"
        "│   │   ├── stt.service.js       # Groq/OpenAI Whisper\n"
        "│   │   ├── llm.service.js       # enrichment\n"
        "│   │   ├── storage.service.js   # S3/R2 upload & signed URL\n"
        "│   │   └── auth.service.js\n"
        "│   ├── models/ or prisma/\n"
        "│   └── utils/\n"
        "├── prisma/schema.prisma         # if using Prisma\n"
        "├── .env.example\n"
        "├── package.json\n"
        "└── README.md",
    )

    # 3. Data model
    add_heading(doc, "3. Data Models (Match Mobile Contracts)", 1)
    add_para(
        doc,
        "These shapes come from src/types/index.ts and the mobile stores. Keep response field names "
        "aligned so the Expo app can switch from AsyncStorage to API with minimal mapping.",
    )

    add_heading(doc, "3.1 User", 2)
    add_code(
        doc,
        "{\n"
        '  "id": "uuid",\n'
        '  "firstName": "string",\n'
        '  "lastName": "string",\n'
        '  "email": "string",\n'
        '  "createdAt": "ISO-8601 datetime"\n'
        "}",
    )
    add_para(doc, "Password rules (from app signup): minimum 8 characters, at least 1 uppercase letter, at least 1 number. Store bcrypt/argon2 hashes only — never plaintext.")

    add_heading(doc, "3.2 Idea", 2)
    add_code(
        doc,
        "{\n"
        '  "id": "uuid",\n'
        '  "userId": "uuid",\n'
        '  "title": "string",\n'
        '  "category": "Movies | Songs | Books | Business | Scripts | Design | Music",\n'
        '  "summary": "string",\n'
        '  "transcript": "string",          // raw STT — source of truth\n'
        '  "aiStory": "string | null",\n'
        '  "audioUri": "https://...signed...", // or storage path resolved by client\n'
        '  "durationSec": 42,\n'
        '  "language": "hi",                 // ISO / Whisper code\n'
        '  "transcriptSource": "live | demo | device",\n'
        '  "favorite": false,\n'
        '  "createdAt": "ISO-8601",\n'
        '  "updatedAt": "ISO-8601"\n'
        "}",
    )

    add_heading(doc, "3.3 AI Enrichment response", 2)
    add_code(
        doc,
        "{\n"
        '  "transcript": "string",\n'
        '  "title": "string",\n'
        '  "category": "Business",\n'
        '  "summary": "string",\n'
        '  "aiStory": "string",\n'
        '  "detectedLanguage": "hi"\n'
        "}",
    )

    add_heading(doc, "3.4 Suggested SQL tables", 2)
    add_para(doc, "users", bold=True)
    add_table(
        doc,
        ["Column", "Type", "Notes"],
        [
            ["id", "UUID PK", ""],
            ["first_name", "TEXT", ""],
            ["last_name", "TEXT", ""],
            ["email", "CITEXT UNIQUE", "case-insensitive"],
            ["password_hash", "TEXT", "bcrypt/argon2"],
            ["created_at", "TIMESTAMPTZ", ""],
            ["updated_at", "TIMESTAMPTZ", ""],
        ],
    )
    add_para(doc, "ideas", bold=True)
    add_table(
        doc,
        ["Column", "Type", "Notes"],
        [
            ["id", "UUID PK", ""],
            ["user_id", "UUID FK → users", "isolate per user"],
            ["title", "TEXT", ""],
            ["category", "TEXT", "English taxonomy"],
            ["summary", "TEXT", ""],
            ["transcript", "TEXT", "never overwrite with aiStory"],
            ["ai_story", "TEXT NULL", ""],
            ["audio_path", "TEXT", "object storage key"],
            ["duration_sec", "INT", ""],
            ["language", "TEXT", "detected spoken language"],
            ["transcript_source", "TEXT", "live | demo | device"],
            ["favorite", "BOOLEAN DEFAULT false", ""],
            ["created_at", "TIMESTAMPTZ", ""],
            ["updated_at", "TIMESTAMPTZ", ""],
        ],
    )
    add_para(doc, "Indexes: (user_id, created_at DESC); full-text on title/summary/transcript; optional index on category.")
    add_para(doc, "user_settings: user_id, language_code, speech_locale, wake_word_enabled")
    add_para(doc, "subscriptions: user_id, name, email, contact, profession, language, plan, started_at, expires_at")
    add_para(doc, "Storage layout: ideas-audio/{user_id}/{idea_id}.m4a — private bucket, signed GET URLs for playback.")

    # 4. API endpoints
    add_heading(doc, "4. Complete API Endpoint Specification", 1)
    add_para(doc, "Base URL example: https://api.thinktap.example")
    add_para(doc, "Auth header (except public auth routes): Authorization: Bearer <access_token>")
    add_para(doc, "Error body (match existing proxy): { \"error\": \"message\" } with appropriate HTTP status.")

    add_heading(doc, "4.1 Health", 2)
    add_table(
        doc,
        ["Method", "Path", "Auth", "Response"],
        [["GET", "/health", "No", '{ "ok": true, "provider": "Groq", "sttModel": "...", "maxUploadBytes": 26214400 }']],
    )

    add_heading(doc, "4.2 Authentication", 2)
    add_table(
        doc,
        ["Method", "Path", "Request body", "Success response"],
        [
            ["POST", "/auth/signup", '{ "firstName", "lastName", "email", "password" }', '{ "user", "accessToken", "refreshToken?" }'],
            ["POST", "/auth/login", '{ "email", "password" }', "same as signup"],
            ["POST", "/auth/logout", "— (Bearer)", '{ "ok": true }'],
            ["POST", "/auth/refresh", '{ "refreshToken" }', '{ "accessToken" }'],
            ["GET", "/auth/me", "Bearer", "User object"],
            ["POST", "/auth/guest", "optional deviceId", "guest session tokens"],
            ["POST", "/auth/migrate-guest", "Bearer + optional guest idea ids", "merged account + ideas"],
        ],
    )
    add_para(doc, "Signup validation: email format; password ≥8 with uppercase + number; firstName/lastName required. Return 409 on duplicate email.")

    add_heading(doc, "4.3 Ideas CRUD", 2)
    add_table(
        doc,
        ["Method", "Path", "Query / body", "Response"],
        [
            ["GET", "/ideas", "?category=&favorite=&limit=&cursor=", '{ "items": Idea[], "nextCursor?" }'],
            ["GET", "/ideas/:id", "—", "Idea"],
            ["POST", "/ideas", "JSON (see below)", "Idea"],
            ["PATCH", "/ideas/:id", "partial Idea fields", "Idea"],
            ["DELETE", "/ideas/:id", "—", '{ "ok": true }'],
            ["GET", "/ideas/search", "?q=", '{ "items": Idea[] }'],
        ],
    )
    add_para(doc, "POST /ideas body (when client already has transcript/enrichment):", bold=True)
    add_code(
        doc,
        "{\n"
        '  "title": "string",\n'
        '  "category": "Business",\n'
        '  "summary": "string",\n'
        '  "transcript": "string",\n'
        '  "aiStory": "string | null",\n'
        '  "durationSec": 42,\n'
        '  "language": "hi",\n'
        '  "transcriptSource": "device",\n'
        '  "audioUploadId": "uuid"\n'
        "}",
    )
    add_para(doc, "Always scope list/get/update/delete by authenticated user_id. Never return another user’s ideas.")

    add_heading(doc, "4.4 Audio upload & playback", 2)
    add_table(
        doc,
        ["Method", "Path", "Body", "Response"],
        [
            ["POST", "/audio/upload", 'multipart field "file"', '{ "uploadId", "audioUrl", "path", "sizeBytes", "mimeType" }'],
            ["GET", "/audio/:id", "—", "302 redirect or { signedUrl, expiresAt }"],
            ["DELETE", "/audio/:id", "—", '{ "ok": true }'],
        ],
    )
    add_para(
        doc,
        "App records mono AAC/m4a (≈16 kHz, 96 kbps) via expo-audio. Detect MIME: .wav→audio/wav, .mp3→audio/mpeg, "
        ".webm→audio/webm, .3gp→audio/3gpp, else audio/mp4 (idea.m4a). Reject uploads over MAX_UPLOAD_BYTES (default 25 MB).",
    )

    add_heading(doc, "4.5 AI processing (existing + extensions)", 2)
    add_table(
        doc,
        ["Method", "Path", "Body", "Response"],
        [
            ["POST", "/api/transcribe", 'multipart "file"', '{ "text", "language" }'],
            ["POST", "/api/enrich", 'multipart "file"', "AiEnrichment (+ detectedLanguage)"],
            ["POST", "/api/enrich-from-transcript", '{ "transcript", "speechLocale?", "language?" }', "AiEnrichment"],
            [
                "POST",
                "/ideas/process",
                'multipart file OR { audioUploadId, transcript?, speechLocale? }',
                '{ "idea": Idea }  // atomic upload + enrich + save',
            ],
        ],
    )
    add_para(
        doc,
        "/api/enrich and /api/transcribe already exist in server/index.js and match the mobile client in "
        "src/services/aiService.ts. Keep these contracts stable. Add enrich-from-transcript to support the "
        "current on-device OS STT path without re-transcribing. Prefer /ideas/process for a one-shot cloud pipeline.",
    )

    add_heading(doc, "4.6 Settings & subscription", 2)
    add_table(
        doc,
        ["Method", "Path", "Body", "Response"],
        [
            ["GET", "/me/settings", "—", '{ "languageCode", "speechLocale", "wakeWordEnabled" }'],
            ["PATCH", "/me/settings", "partial settings", "settings"],
            ["POST", "/subscriptions/trial", '{ "name", "email", "contact", "profession", "language" }', "SubscriptionProfile"],
            ["GET", "/subscriptions/me", "—", "profile or null"],
        ],
    )
    add_para(doc, "profession enum used in app: businessman | working_professional. plan today: free_trial.")

    # 5. Flows
    add_heading(doc, "5. End-to-End Flows Backend Must Support", 1)

    add_heading(doc, "5.1 Current mobile flow (device STT — still primary)", 2)
    add_code(
        doc,
        "Mic → record m4a + parallel OS speech recognition\n"
        "  → pendingRecordingStore { audioUri, durationSec, transcript, speechLocale }\n"
        "  → /processing → local heuristics for title/category/summary\n"
        "  → save Idea locally\n"
        "\n"
        "Backend hook for this path:\n"
        "  1) POST /audio/upload (audio file)\n"
        "  2) POST /api/enrich-from-transcript { transcript }\n"
        "  3) POST /ideas { ...enrichment, audioUploadId, durationSec, ... }",
    )

    add_heading(doc, "5.2 Target cloud AI flow", 2)
    add_code(
        doc,
        "Mic → record m4a\n"
        "  → POST /ideas/process (multipart file + durationSec)\n"
        "       server: store audio → Whisper STT (auto language) → LLM enrich → INSERT idea\n"
        "  → return Idea JSON with signed audioUri\n"
        "  → app navigates to /idea/{id}",
    )

    add_heading(doc, "5.3 Auth + sync", 2)
    add_code(
        doc,
        "App launch → refresh token → GET /ideas\n"
        "Offline queue (optional later) → sync when online\n"
        "Guest record → signup → POST /auth/migrate-guest (rewrite userId on ideas)",
    )

    add_heading(doc, "5.4 Processing stages (UI contract)", 2)
    add_para(
        doc,
        "Mobile processing screen shows stages: uploading → transcribing → extracting → summarizing → saving → done | error. "
        "If you stream progress later, map server events to these names. Failures must not silently drop user audio when avoidable.",
    )

    # 6. Step by step
    add_heading(doc, "6. Step-by-Step Implementation Plan for Backend Developer", 1)

    add_heading(doc, "Step 1 — Project setup", 2)
    add_numbered(doc, "Create backend/ (or extend existing server/).")
    add_numbered(doc, "npm init; install express, cors, dotenv, multer, bcrypt (or argon2), jsonwebtoken, zod, helmet, express-rate-limit.")
    add_numbered(doc, "Add Prisma/Drizzle + PostgreSQL client, and AWS SDK or @supabase/supabase-js for storage.")
    add_numbered(doc, "Copy env template from server/.env.example and expand (see Step 2).")
    add_numbered(doc, "Confirm GET /health works on PORT 8787 (or your chosen port).")

    add_heading(doc, "Step 2 — Environment & secrets", 2)
    add_code(
        doc,
        "PORT=8787\n"
        "NODE_ENV=development\n"
        "DATABASE_URL=postgresql://...\n"
        "JWT_ACCESS_SECRET=...\n"
        "JWT_REFRESH_SECRET=...\n"
        "JWT_ACCESS_TTL=15m\n"
        "JWT_REFRESH_TTL=30d\n"
        "GROQ_API_KEY=gsk_...          # preferred for MVP cost\n"
        "# OPENAI_API_KEY=sk_...\n"
        "STT_MODEL=whisper-large-v3\n"
        "CHAT_MODEL=llama-3.3-70b-versatile\n"
        "MAX_UPLOAD_BYTES=26214400\n"
        "STORAGE_BUCKET=ideas-audio\n"
        "STORAGE_ENDPOINT=...\n"
        "STORAGE_ACCESS_KEY=...\n"
        "STORAGE_SECRET_KEY=...\n"
        "CORS_ORIGINS=*",
    )
    add_para(doc, "Never commit real keys. Rotate any key that was ever under EXPO_PUBLIC_*.")

    add_heading(doc, "Step 3 — Database schema & migrations", 2)
    add_numbered(doc, "Create users, ideas, user_settings, subscriptions tables as in Section 3.4.")
    add_numbered(doc, "Add FK user_id on ideas with ON DELETE CASCADE (or soft-delete policy).")
    add_numbered(doc, "Add indexes for list + search.")
    add_numbered(doc, "Seed a test user for Postman/Insomnia.")

    add_heading(doc, "Step 4 — Auth module", 2)
    add_numbered(doc, "Implement signup: validate → hash password → insert user → issue tokens.")
    add_numbered(doc, "Implement login: verify hash → issue tokens.")
    add_numbered(doc, "Middleware: read Bearer token, verify JWT, attach req.user = { id, email }.")
    add_numbered(doc, "Implement /auth/me, logout (optional token blacklist), refresh.")
    add_numbered(doc, "Optional: guest session + migrate-guest for record-before-login.")

    add_heading(doc, "Step 5 — Keep & harden existing AI proxy", 2)
    add_numbered(doc, "Preserve POST /api/transcribe and POST /api/enrich contracts exactly.")
    add_numbered(doc, "STT: multipart to provider /audio/transcriptions; omit language for auto-detect.")
    add_numbered(doc, "Use transcription prompt that preserves Devanagari for Hindi/Marathi (see server/index.js AUTO_PROMPT).")
    add_numbered(doc, "LLM: JSON object with title (≤8 words), English category from fixed list, summary 1–2 sentences, aiStory, transcript copied unchanged.")
    add_numbered(doc, "Map errors: 400 missing file, 413 too large, 422 no speech, 429 rate limit, 502/503 upstream.")
    add_numbered(doc, "Add rate limiting per IP/user before production.")

    add_heading(doc, "Step 6 — Audio storage service", 2)
    add_numbered(doc, "Implement POST /audio/upload with multer memory/disk temp.")
    add_numbered(doc, "Upload to private bucket path ideas-audio/{userId}/{uuid}.m4a.")
    add_numbered(doc, "Return uploadId + path; generate short-lived signed URLs for GET playback.")
    add_numbered(doc, "On idea delete, delete object from storage.")
    add_numbered(doc, "Delete temp files after processing.")

    add_heading(doc, "Step 7 — Ideas CRUD + search", 2)
    add_numbered(doc, "Implement GET/POST/PATCH/DELETE /ideas with ownership checks.")
    add_numbered(doc, "Map DB columns → mobile Idea JSON (camelCase). Resolve audioUri via signed URL.")
    add_numbered(doc, "Implement GET /ideas/search?q= with ILIKE or Postgres full-text on title, summary, transcript, category.")
    add_numbered(doc, "Support category filter chips and favorite filter.")

    add_heading(doc, "Step 8 — Atomic process endpoint", 2)
    add_numbered(doc, "Implement POST /ideas/process: validate auth → store audio → STT → enrich → insert idea → return Idea.")
    add_numbered(doc, "Also implement POST /api/enrich-from-transcript for device-STT clients.")
    add_numbered(doc, "Ensure raw transcript is never overwritten by polished aiStory.")

    add_heading(doc, "Step 9 — Settings & subscription (optional MVP+)", 2)
    add_numbered(doc, "GET/PATCH /me/settings for languageCode, speechLocale, wakeWordEnabled.")
    add_numbered(doc, "POST /subscriptions/trial to persist free-trial profile (currently local-only in app).")

    add_heading(doc, "Step 10 — Security & production hardening", 2)
    add_numbered(doc, "Put TLS reverse proxy (Nginx/Caddy) in front.")
    add_numbered(doc, "CORS allowlist for your app / web admin only.")
    add_numbered(doc, "helmet + body size limits + upload size limits.")
    add_numbered(doc, "Rate-limit /api/enrich, /ideas/process, /auth/login.")
    add_numbered(doc, "Structured logging; never log full transcripts in shared logs if privacy-sensitive.")
    add_numbered(doc, "Plan chunking for >25 MB audio (~10 min segments + prior-segment prompt) — documented but not yet implemented.")

    add_heading(doc, "Step 11 — Mobile integration checklist", 2)
    add_numbered(doc, "Set EXPO_PUBLIC_API_URL to your LAN IP:port for device testing (not localhost on physical devices).")
    add_numbered(doc, "Point authStore login/signup at /auth/* instead of AsyncStorage users.")
    add_numbered(doc, "Point ideasStore load/add/update/delete at /ideas.")
    add_numbered(doc, "Wire processing.tsx to /ideas/process or upload + enrich-from-transcript + POST /ideas.")
    add_numbered(doc, "AudioPlayer uses signed HTTPS audioUri instead of file://.")
    add_numbered(doc, "Remove EXPO_PUBLIC_GROQ_API_KEY from production builds.")

    add_heading(doc, "Step 12 — Test plan", 2)
    add_bullet(doc, "GET /health returns provider when key set.")
    add_bullet(doc, "Signup → login → /auth/me round-trip.")
    add_bullet(doc, "Upload short m4a → /api/enrich returns title/category/summary/transcript/detectedLanguage.")
    add_bullet(doc, "Hindi/Marathi sample preserves Devanagari (not romanized).")
    add_bullet(doc, "POST /ideas/process creates Idea owned by user; GET /ideas lists it.")
    add_bullet(doc, "Search finds by transcript substring.")
    add_bullet(doc, "PATCH favorite; DELETE removes DB row + storage object.")
    add_bullet(doc, "Unauthorized access to another user’s idea → 403/404.")
    add_bullet(doc, "Oversized file → 413; empty speech → 422; missing key → 503.")
    add_bullet(doc, "End-to-end from Expo app on Android device using LAN API URL.")

    # 7. LLM prompt
    add_heading(doc, "7. LLM Enrichment Prompt Contract", 1)
    add_para(doc, "Use a system prompt equivalent to the existing server (keep category list in English):")
    add_code(
        doc,
        "You organize voice ideas for Think Tap.\n"
        "Detected spoken language code: {languageCode}.\n"
        "The transcript below is already in the correct script — copy it into \"transcript\" unchanged.\n"
        "Write title, summary, and aiStory in the same language/script as the transcript.\n"
        "Keep category in English from: Movies, Songs, Books, Business, Scripts, Design, Music.\n"
        "Title: max 8 words. Summary: 1-2 sentences from the transcript only.\n"
        "Return JSON: { \"transcript\", \"title\", \"category\", \"summary\", \"aiStory\" }",
    )
    add_para(doc, "Fallback: if LLM fails, derive title from first sentence (≤8 words), category Business, summary truncated transcript — same as localEnrich in server/index.js.")

    # 8. Status codes
    add_heading(doc, "8. HTTP Status Codes", 1)
    add_table(
        doc,
        ["Status", "When"],
        [
            ["200 / 201", "Success"],
            ["400", "Missing/invalid body or multipart file"],
            ["401", "Missing/invalid token"],
            ["403", "Valid token but not owner"],
            ["404", "Idea/audio/user not found"],
            ["409", "Duplicate email on signup"],
            ["413", "Audio exceeds MAX_UPLOAD_BYTES (~25 MB)"],
            ["422", "No speech detected"],
            ["429", "Rate limited (app or upstream)"],
            ["500", "Unexpected server error"],
            ["502", "Upstream STT/LLM bad response"],
            ["503", "Missing provider API key / service unavailable"],
        ],
    )

    # 9. What NOT to build
    add_heading(doc, "9. Out of Scope for Backend (Do Not Build)", 1)
    add_bullet(doc, "Wake word detection APIs — on-device Android/iOS only.")
    add_bullet(doc, "Client-side Groq/OpenAI keys in production.")
    add_bullet(doc, "Vector/semantic search (Phase 2+).")
    add_bullet(doc, "AES-256 / .onetap encryption format (Phase 4).")
    add_bullet(doc, "Stripe billing (unless client expands free_trial).")

    # 10. Priority
    add_heading(doc, "10. Implementation Priority Order", 1)
    add_numbered(doc, "Keep /health, /api/transcribe, /api/enrich working and secured.")
    add_numbered(doc, "Auth + JWT middleware.")
    add_numbered(doc, "Audio storage upload + signed URLs.")
    add_numbered(doc, "Ideas CRUD + search.")
    add_numbered(doc, "/ideas/process atomic pipeline.")
    add_numbered(doc, "/api/enrich-from-transcript.")
    add_numbered(doc, "Settings / subscription sync.")
    add_numbered(doc, "Rate limits, TLS, monitoring, long-audio chunking.")

    # 11. Reference files
    add_heading(doc, "11. Source-of-Truth Files in This Repo", 1)
    add_table(
        doc,
        ["Path", "Why it matters"],
        [
            ["src/types/index.ts", "Canonical Idea / User / AiEnrichment types"],
            ["src/services/aiService.ts", "Client enrich API contract"],
            ["server/index.js", "Existing STT/LLM proxy implementation"],
            ["server/.env.example", "Server secrets template"],
            ["server/README.md", "Proxy setup notes"],
            ["app/processing.tsx", "Create-idea pipeline stages"],
            ["src/store/authStore.ts", "Auth session shape to replace"],
            ["src/store/ideasStore.ts", "Ideas persistence to replace"],
            [".env.example", "EXPO_PUBLIC_API_URL wiring"],
            ["README.md", "Product phases + recommended Supabase stack"],
        ],
    )

    # 12. Quick start
    add_heading(doc, "12. Quick Start (Existing Proxy Today)", 1)
    add_code(
        doc,
        "cd server\n"
        "cp .env.example .env\n"
        "# set GROQ_API_KEY or OPENAI_API_KEY\n"
        "npm install\n"
        "npm start\n"
        "\n"
        "# In Expo .env:\n"
        "EXPO_PUBLIC_API_URL=http://YOUR_LAN_IP:8787",
    )
    add_para(
        doc,
        "Cost ballpark: Whisper / gpt-4o-transcribe class ≈ $0.004–0.006 per audio minute. Prefer Groq for MVP cost; OpenAI for higher accuracy.",
    )

    add_heading(doc, "13. Definition of Done", 1)
    add_bullet(doc, "All Section 4 endpoints implemented with ownership checks.")
    add_bullet(doc, "Provider keys server-only.")
    add_bullet(doc, "Mobile can sign up, record, process via API, list/search/play/delete ideas.")
    add_bullet(doc, "Audio stored privately; playback via signed URL.")
    add_bullet(doc, "Error codes match Section 8.")
    add_bullet(doc, "No secrets in git; health check green in staging.")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(
        "\n— End of document —\n"
        "Think Tap Backend API Implementation Guide (Node.js)\n"
        "Derived from the Think Tap Expo codebase and existing server/ STT proxy."
    )
    set_run_font(r, size=10, color=RGBColor(0x6B, 0x72, 0x80))

    out = Path(__file__).resolve().parent / "ThinkTap_Backend_API_Implementation_Guide.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    build()
